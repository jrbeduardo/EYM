import numpy as np
import pandas as pd
import matplotlib
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Cm
import matplotlib.mathtext as mathtext
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


titulo=r'Gráfica $\mu$'
ejeX=r'$x\left[C\right]$'
ejeY=r'$x\left[m^{2}\right]$'

data = {'x': [29, 9, 10, 38, 16, 26],
            'y': [65, 7, 8, 76, 23, 56]}

df = pd.DataFrame(data=data)
x = df['x']
y = df['y']
model = np.polyfit(x, y, 1)



x_data = np.arange(0, 51, 1)
y_data = model[0] * x_data + model[1]

pendiente=round(model[0],4)

ordenada=round(model[1],4)
if(ordenada>=0):
	modelo=r'$y\left[ m \right] = {}\left[ N \right]\cdot x\left[ \frac{{m}}{{N}} \right]+{}\left[ m \right]$'.format(pendiente,ordenada)
else: 
	modelo=r'$y\left[ m \right] = {}\left[ N \right]\cdot x\left[ \frac{{m}}{{N}} \right]{}\left[ m \right]$'.format(pendiente,ordenada)



fig, ax = plt.subplots(figsize=(7,7))
ax = plt.gca() 
ax.spines['right'].set_color('none')
ax.spines['top'].set_color('none')
ax.xaxis.set_ticks_position('bottom')
ax.spines['bottom'].set_position(('data',0))
ax.yaxis.set_ticks_position('left')
ax.spines['left'].set_position(('data',0))

plt.scatter(x, y)

ax.plot(x_data,y_data)

ax.plot(x_data, y_data, c = 'r')

ax.set_title(titulo, size=24)
ax.set_xlabel(ejeX, size=24)
ax.set_ylabel(ejeY, size=24)

ax.grid(True, which='both', linestyle='--')
#plt.savefig("grafica.svg", bbox_inches = 'tight', pad_inches = .2)

plt.savefig("grafica.png",pad_inches = .2)


document = Document()

sections = document.sections
for section in sections:
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

t = document.add_table(df.shape[0]+1, df.shape[1])

# add the header rows.
for j in range(df.shape[-1]):
    t.cell(0,j).text = df.columns[j]

# add the rest of the data frame
for i in range(df.shape[0]):
    for j in range(df.shape[-1]):
        t.cell(i+1,j).text = str(df.values[i,j])

t.style = 'Table Grid'
t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER     




inline_obj= document.add_picture("grafica.png") #return Inline object
inline_obj.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER     


parser=mathtext.MathTextParser( 'bitmap' )
offset=parser.to_png("modelo.png", modelo, fontsize=12)
inline_obj= document.add_picture("modelo.png") #return Inline object
inline_obj.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER     



document.save('grafica.docx')